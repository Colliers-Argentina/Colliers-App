#!/usr/bin/env python3
"""
Generador de Guías Rápidas de Usuario de Colliers Nexus (formato simple).

Política (ver CLAUDE.md > "Documentación de formularios"):
  - TODOS los formularios del sistema tienen su guía rápida (no solo los nuevos).
  - Guía breve para el usuario final: Objetivo, Campos ("Qué ingresar"), Al finalizar.
  - Sin capturas, sin diagramas, sin arquitectura, sin reglas técnicas, sin código.
  - Formato Word (.docx) + PDF, generados desde el mismo contenido.
  - Cuando un formulario cambia, se ACTUALIZA su guía (no se crea una nueva).

Uso: definir una Guia(...) y llamar guia.build(dir, "Guia_<Form>").
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, Image, HRFlowable, ListFlowable, ListItem)
from reportlab.lib.enums import TA_CENTER

NAVY = RGBColor(0x00, 0x28, 0x5A)
BLUE = RGBColor(0x1C, 0x54, 0xF4)
GREY = RGBColor(0x56, 0x64, 0x8F)
DARK = RGBColor(0x22, 0x2A, 0x40)

R_NAVY = colors.HexColor("#00285A")
R_BLUE = colors.HexColor("#1C54F4")
R_GREY = colors.HexColor("#56648F")
R_LINE = colors.HexColor("#DBE5FF")

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(HERE, "..", ".."))
LOGO = os.path.join(REPO, "colliers-logo-white.png")


def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


class Guia:
    def __init__(self, nombre, version, fecha, objetivo, campos, al_finalizar):
        self.nombre = nombre            # p.ej. "Cuentas"
        self.version = version
        self.fecha = fecha
        self.objetivo = objetivo        # str (2-3 líneas)
        self.campos = campos            # list[(nombre, que_ingresar)]
        self.al_finalizar = al_finalizar  # list[str]

    @property
    def titulo(self):
        return f"Manual de Usuario - {self.nombre}"

    # ───────── DOCX ─────────
    def to_docx(self, path):
        doc = Document()
        st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11); st.font.color.rgb = DARK

        band = doc.add_table(rows=1, cols=1)
        c = band.rows[0].cells[0]; _shade(c, "00285A")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(LOGO):
            c.paragraphs[0].add_run().add_picture(LOGO, width=Inches(1.6))

        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(self.titulo); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
        meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = meta.add_run(f"Versión {self.version}   ·   {self.fecha}")
        r.font.size = Pt(10.5); r.font.color.rgb = GREY
        doc.add_paragraph()

        self._docx_h(doc, "Objetivo")
        doc.add_paragraph(self.objetivo)
        doc.add_paragraph()

        self._docx_h(doc, "Campos")
        for i, (nombre, que) in enumerate(self.campos):
            p = doc.add_paragraph(); rn = p.add_run(nombre); rn.bold = True; rn.font.size = Pt(11.5); rn.font.color.rgb = NAVY
            lbl = doc.add_paragraph(); rl = lbl.add_run("Qué ingresar:"); rl.bold = True; rl.font.color.rgb = BLUE; rl.font.size = Pt(10.5)
            doc.add_paragraph(que)
            if i < len(self.campos) - 1:
                sep = doc.add_paragraph(); sr = sep.add_run("─" * 46); sr.font.color.rgb = RGBColor(0xC9, 0xD6, 0xF0)

        doc.add_paragraph()
        self._docx_h(doc, "Al finalizar")
        intro = doc.add_paragraph(); ir = intro.add_run("Al guardar:"); ir.bold = True
        for b in self.al_finalizar:
            doc.add_paragraph(b, style="List Bullet")

        doc.save(path); return path

    def _docx_h(self, doc, text):
        p = doc.add_paragraph(); r = p.add_run(text)
        r.bold = True; r.font.size = Pt(14); r.font.color.rgb = BLUE

    # ───────── PDF ─────────
    def to_pdf(self, path):
        base = ParagraphStyle("b", fontName="Helvetica", fontSize=11, leading=15, textColor=DARK_R())
        h = ParagraphStyle("h", parent=base, fontName="Helvetica-Bold", fontSize=14,
                           textColor=R_BLUE, spaceBefore=8, spaceAfter=5)
        field = ParagraphStyle("f", parent=base, fontName="Helvetica-Bold", fontSize=11.5,
                               textColor=R_NAVY, spaceBefore=6, spaceAfter=1)
        qlabel = ParagraphStyle("q", parent=base, fontName="Helvetica-Bold", fontSize=10.5,
                                textColor=R_BLUE, spaceAfter=0)
        story = []

        if os.path.exists(LOGO):
            story.append(Image(LOGO, width=3.6 * cm, height=3.6 * cm * 196 / 270, hAlign="CENTER"))
            story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph(self.titulo, ParagraphStyle("t", parent=base, fontName="Helvetica-Bold",
                     fontSize=20, textColor=R_NAVY, alignment=TA_CENTER)))
        story.append(Paragraph(f"Versión {self.version}  ·  {self.fecha}",
                     ParagraphStyle("m", parent=base, fontSize=10.5, textColor=R_GREY, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.5 * cm))

        def esc(s):
            return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        story.append(Paragraph("Objetivo", h))
        story.append(Paragraph(esc(self.objetivo), base))
        story.append(Spacer(1, 0.3 * cm))

        story.append(Paragraph("Campos", h))
        for i, (nombre, que) in enumerate(self.campos):
            story.append(Paragraph(esc(nombre), field))
            story.append(Paragraph("Qué ingresar:", qlabel))
            story.append(Paragraph(esc(que), base))
            if i < len(self.campos) - 1:
                story.append(Spacer(1, 0.15 * cm))
                story.append(HRFlowable(width="100%", thickness=0.5, color=R_LINE))

        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("Al finalizar", h))
        story.append(Paragraph("<b>Al guardar:</b>", base))
        story.append(ListFlowable([ListItem(Paragraph(esc(b), base)) for b in self.al_finalizar],
                     bulletType="bullet", start="•", leftIndent=14))

        SimpleDocTemplate(path, pagesize=A4, leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                          topMargin=1.6 * cm, bottomMargin=1.6 * cm,
                          title=self.titulo, author="Colliers Nexus").build(story)
        return path

    def build(self, basedir, basename):
        d = self.to_docx(os.path.join(basedir, basename + ".docx"))
        p = self.to_pdf(os.path.join(basedir, basename + ".pdf"))
        return d, p


def DARK_R():
    return colors.HexColor("#222A40")
