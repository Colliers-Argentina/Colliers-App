#!/usr/bin/env python3
"""
Generador base de Manuales de Usuario de Colliers Nexus.

Los manuales se describen como una LISTA DE BLOQUES (estructura de datos neutral)
y se renderizan a DOCX (python-docx) y a PDF (reportlab) desde esa misma fuente,
de modo que ambos formatos queden siempre sincronizados.

Reglas del proyecto (ver CLAUDE.md > "Manuales de Usuario"):
  - Cada formulario nuevo o cambio relevante lleva su manual DOCX + PDF.
  - Si el manual ya existe, se ACTUALIZA (nueva fila en Historial de cambios),
    no se duplica.
  - La documentación se mantiene sincronizada con la versión actual de Nexus.

Tipos de bloque (dicts):
  {"t":"h1","n":2,"text":"..."}          Título de sección numerada
  {"t":"h2","text":"..."}                 Subtítulo
  {"t":"p","text":"..."}                  Párrafo
  {"t":"bullets","items":[...]}           Lista con viñetas
  {"t":"steps","items":[...]}             Lista numerada
  {"t":"shot","n":1,"caption":"..."}      Marcador de captura de pantalla
  {"t":"table","headers":[...],"rows":[[...]],"widths":[...]}
  {"t":"faq","items":[(pregunta,respuesta),...]}
  {"t":"spacer"}
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, Image, PageBreak, ListFlowable, ListItem)
from reportlab.lib.enums import TA_CENTER

NAVY = RGBColor(0x00, 0x28, 0x5A)
BLUE = RGBColor(0x1C, 0x54, 0xF4)
GREY = RGBColor(0x56, 0x64, 0x8F)
LIGHT = RGBColor(0x88, 0x98, 0xC0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

R_NAVY = colors.HexColor("#00285A")
R_BLUE = colors.HexColor("#1C54F4")
R_GREY = colors.HexColor("#56648F")
R_LIGHT = colors.HexColor("#8898C0")
R_SHOT = colors.HexColor("#EEF3FF")

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(HERE, "..", ".."))
LOGO = os.path.join(REPO, "colliers-logo-white.png")


# ───────────────────────── helpers DOCX ─────────────────────────
def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


class Manual:
    """Contiene metadatos + lista de bloques. Renderiza a DOCX y PDF."""

    def __init__(self, modulo, subtitulo, version, fecha):
        self.modulo = modulo
        self.subtitulo = subtitulo
        self.version = version
        self.fecha = fecha
        self.blocks = []

    # API de contenido
    def h1(self, n, text): self.blocks.append({"t": "h1", "n": n, "text": text}); return self
    def h2(self, text): self.blocks.append({"t": "h2", "text": text}); return self
    def p(self, text): self.blocks.append({"t": "p", "text": text}); return self
    def bullets(self, items): self.blocks.append({"t": "bullets", "items": list(items)}); return self
    def steps(self, items): self.blocks.append({"t": "steps", "items": list(items)}); return self
    def shot(self, n, caption): self.blocks.append({"t": "shot", "n": n, "caption": caption}); return self
    def table(self, headers, rows, widths=None):
        self.blocks.append({"t": "table", "headers": headers, "rows": rows, "widths": widths}); return self
    def faq(self, items): self.blocks.append({"t": "faq", "items": items}); return self
    def spacer(self): self.blocks.append({"t": "spacer"}); return self

    # ───────────────────── render DOCX ─────────────────────
    def to_docx(self, path):
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name = "Calibri"; st.font.size = Pt(10.5)
        st.font.color.rgb = RGBColor(0x22, 0x2A, 0x40)

        band = doc.add_table(rows=1, cols=1); band.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = band.rows[0].cells[0]; _shade(c, "00285A")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(LOGO):
            c.paragraphs[0].add_run().add_picture(LOGO, width=Inches(2.1))
        else:
            _cell(c, "COLLIERS", True, WHITE, 26, WD_ALIGN_PARAGRAPH.CENTER)
        pad = c.add_paragraph(); pad.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pad.add_run("NEXUS"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = WHITE

        doc.add_paragraph()
        for text, size, color, bold in [
            ("Manual de Usuario", 13, GREY, False),
            (self.modulo, 24, NAVY, True),
            (self.subtitulo, 11, GREY, False)]:
            pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pr.add_run(text); rr.font.size = Pt(size); rr.font.color.rgb = color; rr.bold = bold
        doc.add_paragraph()
        meta = doc.add_table(rows=2, cols=2); meta.alignment = WD_TABLE_ALIGNMENT.CENTER; meta.style = "Table Grid"
        _cell(meta.rows[0].cells[0], "Versión", True, NAVY); _cell(meta.rows[0].cells[1], self.version)
        _cell(meta.rows[1].cells[0], "Fecha", True, NAVY); _cell(meta.rows[1].cells[1], self.fecha)
        nt = doc.add_paragraph(); nt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = nt.add_run("Colliers Argentina · Documento de uso interno"); rr.font.size = Pt(8); rr.font.color.rgb = LIGHT
        doc.add_page_break()

        for b in self.blocks:
            t = b["t"]
            if t == "h1":
                pr = doc.add_paragraph(); rr = pr.add_run(f'{b["n"]}. {b["text"]}')
                rr.bold = True; rr.font.size = Pt(15); rr.font.color.rgb = NAVY
            elif t == "h2":
                pr = doc.add_paragraph(); rr = pr.add_run(b["text"])
                rr.bold = True; rr.font.size = Pt(11.5); rr.font.color.rgb = BLUE
            elif t == "p":
                doc.add_paragraph(b["text"])
            elif t == "bullets":
                for it in b["items"]:
                    doc.add_paragraph(it, style="List Bullet")
            elif t == "steps":
                for it in b["items"]:
                    doc.add_paragraph(it, style="List Number")
            elif t == "shot":
                tbl = doc.add_table(rows=2, cols=1); tbl.style = "Table Grid"
                top = tbl.rows[0].cells[0]; _shade(top, "EEF3FF")
                _cell(top, f'[ CAPTURA {b["n"]} ]', True, BLUE, 11, WD_ALIGN_PARAGRAPH.CENTER)
                for _ in range(3):
                    top.add_paragraph()
                _cell(tbl.rows[1].cells[0], f'Figura {b["n"]}. {b["caption"]}', False, GREY, 9, WD_ALIGN_PARAGRAPH.CENTER)
                doc.add_paragraph()
            elif t == "table":
                headers = b["headers"]; tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Table Grid"
                for i, hd in enumerate(headers):
                    cc = tb.rows[0].cells[i]; _shade(cc, "00285A"); _cell(cc, hd, True, WHITE, 9.5)
                for row in b["rows"]:
                    cells = tb.add_row().cells
                    for i, v in enumerate(row):
                        _cell(cells[i], str(v), size=9.5)
                if b.get("widths"):
                    for row in tb.rows:
                        for i, w in enumerate(b["widths"]):
                            row.cells[i].width = Inches(w)
            elif t == "faq":
                for q, a in b["items"]:
                    pr = doc.add_paragraph(); rr = pr.add_run("P: " + q); rr.bold = True
                    doc.add_paragraph("R: " + a)
            elif t == "spacer":
                doc.add_paragraph()
        doc.save(path)
        return path

    # ───────────────────── render PDF ─────────────────────
    def to_pdf(self, path):
        styles = getSampleStyleSheet()
        body = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                              fontSize=10, leading=14, textColor=colors.HexColor("#222A40"))
        h1 = ParagraphStyle("h1", parent=body, fontName="Helvetica-Bold", fontSize=15,
                            textColor=R_NAVY, spaceBefore=10, spaceAfter=5)
        h2 = ParagraphStyle("h2", parent=body, fontName="Helvetica-Bold", fontSize=11.5,
                            textColor=R_BLUE, spaceBefore=6, spaceAfter=3)
        cap = ParagraphStyle("cap", parent=body, fontSize=8.5, textColor=R_GREY, alignment=TA_CENTER)
        story = []

        # Portada
        story.append(Spacer(1, 1.5 * cm))
        if os.path.exists(LOGO):
            story.append(Image(LOGO, width=5.2 * cm, height=5.2 * cm * 196 / 270,
                               hAlign="CENTER"))
        # banda navy detrás del logo es difícil en flow; usamos título navy
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("NEXUS", ParagraphStyle("nx", parent=body, fontName="Helvetica-Bold",
                     fontSize=16, textColor=R_NAVY, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph("Manual de Usuario", ParagraphStyle("mu", parent=body, fontSize=13,
                     textColor=R_GREY, alignment=TA_CENTER)))
        story.append(Paragraph(self.modulo, ParagraphStyle("mod", parent=body, fontName="Helvetica-Bold",
                     fontSize=24, textColor=R_NAVY, alignment=TA_CENTER, spaceBefore=6, spaceAfter=6)))
        story.append(Paragraph(self.subtitulo, ParagraphStyle("sub", parent=body, fontSize=11,
                     textColor=R_GREY, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.8 * cm))
        meta = Table([["Versión", self.version], ["Fecha", self.fecha]], colWidths=[3 * cm, 6 * cm], hAlign="CENTER")
        meta.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, R_LIGHT),
            ("TEXTCOLOR", (0, 0), (0, -1), R_NAVY),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
        story.append(meta)
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph("Colliers Argentina · Documento de uso interno",
                     ParagraphStyle("ft", parent=body, fontSize=8, textColor=R_LIGHT, alignment=TA_CENTER)))
        story.append(PageBreak())

        def esc(s):
            return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        for b in self.blocks:
            t = b["t"]
            if t == "h1":
                story.append(Paragraph(f'{b["n"]}. {esc(b["text"])}', h1))
            elif t == "h2":
                story.append(Paragraph(esc(b["text"]), h2))
            elif t == "p":
                story.append(Paragraph(esc(b["text"]), body))
            elif t == "bullets":
                story.append(ListFlowable([ListItem(Paragraph(esc(i), body)) for i in b["items"]],
                             bulletType="bullet", start="•", leftIndent=14))
            elif t == "steps":
                story.append(ListFlowable([ListItem(Paragraph(esc(i), body)) for i in b["items"]],
                             bulletType="1", leftIndent=14))
            elif t == "shot":
                ph = Table([[Paragraph(f'<b>[ CAPTURA {b["n"]} ]</b>',
                            ParagraphStyle("s", parent=body, textColor=R_BLUE, alignment=TA_CENTER, fontSize=11))]],
                           colWidths=[16 * cm], rowHeights=[2.6 * cm])
                ph.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), R_SHOT),
                            ("BOX", (0, 0), (-1, -1), 0.5, R_BLUE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
                story.append(ph)
                story.append(Paragraph(f'Figura {b["n"]}. {esc(b["caption"])}', cap))
                story.append(Spacer(1, 0.3 * cm))
            elif t == "table":
                headers = b["headers"]
                data = [[Paragraph(f"<b>{esc(h)}</b>", ParagraphStyle("th", parent=body, textColor=colors.white, fontSize=9)) for h in headers]]
                for row in b["rows"]:
                    data.append([Paragraph(esc(v), ParagraphStyle("td", parent=body, fontSize=9)) for v in row])
                widths = b.get("widths")
                col_w = [w * cm * 2.54 / 2.54 for w in widths] if widths else None
                if widths:
                    total = 16.0
                    scaled = [w / sum(widths) * total for w in widths]
                    col_w = [w * cm for w in scaled]
                tb = Table(data, colWidths=col_w, hAlign="LEFT", repeatRows=1)
                tb.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), R_NAVY),
                    ("GRID", (0, 0), (-1, -1), 0.5, R_LIGHT),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5)]))
                story.append(tb)
            elif t == "faq":
                for q, a in b["items"]:
                    story.append(Paragraph(f"<b>P: {esc(q)}</b>", body))
                    story.append(Paragraph(f"R: {esc(a)}", body))
                    story.append(Spacer(1, 0.15 * cm))
            elif t == "spacer":
                story.append(Spacer(1, 0.35 * cm))

        SimpleDocTemplate(path, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
                          topMargin=1.8 * cm, bottomMargin=1.8 * cm,
                          title=f"Manual {self.modulo}", author="Colliers Nexus").build(story)
        return path

    def build(self, basedir, basename):
        docx_path = os.path.join(basedir, basename + ".docx")
        pdf_path = os.path.join(basedir, basename + ".pdf")
        self.to_docx(docx_path)
        self.to_pdf(pdf_path)
        return docx_path, pdf_path
